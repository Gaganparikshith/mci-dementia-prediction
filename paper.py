"""
build_paper_with_figures.py
===========================
Generates the complete IEEE Access paper as a Word document
with all figures embedded from the saved plot directories.

Run this script on your Windows machine:
    .\.venv\Scripts\python.exe build_paper_with_figures.py

Output:
    Reports\paper_drafts\IEEE_Access_Final_Draft.docx
"""

import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ─────────────────────────────────────────────────────────────────────
PROJECT_ROOT = r"C:\Users\ASUS\PycharmProjects\PythonProject1"
DATA_ROOT    = r"C:\Users\ASUS\Desktop\Research Resources\DementiaResearch"
OUT_DIR      = os.path.join(PROJECT_ROOT, "Reports", "paper_drafts")
OUT_FILE     = os.path.join(OUT_DIR, "IEEE_Access_Final_Draft.docx")
os.makedirs(OUT_DIR, exist_ok=True)

PLOTS = os.path.join(DATA_ROOT, "plots")
OUTS  = os.path.join(PROJECT_ROOT, "Outputs")

# Figure paths — tries multiple candidate locations
def find_fig(*candidates):
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return None

FIGURES = {
    "fig1_roc_p1": find_fig(
        os.path.join(PLOTS, "roc", "roc_curve_xgb.png"),
        os.path.join(OUTS,  "Phase_1_Clinical", "figures", "roc_curve_xgb.png"),
        os.path.join(DATA_ROOT, "xgb_roc_curve.png"),
    ),
    "fig2_shap_beeswarm": find_fig(
        os.path.join(OUTS,  "Phase_1_Clinical", "figures", "shap_xgb_beeswarm.png"),
        os.path.join(PLOTS, "shap", "shap_xgb_beeswarm.png"),
    ),
    "fig3_shap_bar": find_fig(
        os.path.join(OUTS,  "Phase_1_Clinical", "figures", "shap_bar_xgboost.png"),
        os.path.join(PLOTS, "shap", "shap_bar_xgboost.png"),
    ),
    "fig4_shap_dep_adas": find_fig(
        os.path.join(OUTS,  "Phase_1_Clinical", "figures", "shap_xgb_dependence_ADAS-Cog_13.png"),
        os.path.join(PLOTS, "shap", "shap_xgb_dependence_ADAS-Cog_13.png"),
    ),
    "fig5_phase15_shap": find_fig(
        os.path.join(PLOTS, "phase15", "phase15_shap_grouped_bar.png"),
        os.path.join(PLOTS, "phase15_shap_grouped_bar.png"),
        os.path.join(DATA_ROOT, "phase15_shap_grouped_bar.png"),
        os.path.join(OUTS, "Phase_1_5_Clinical_MRI_Volumes", "figures", "phase15_shap_grouped_bar.png"),
    ),
    "fig6_cnn_curves": find_fig(
        os.path.join(OUTS,  "Phase_2_Raw_MRI_CNN", "figures", "cnn2d_v3_training_curves.png"),
        os.path.join(PLOTS, "cnn2d_v3", "cnn2d_v3_training_curves.png"),
    ),
    "fig7_cnn_compare": find_fig(
        os.path.join(OUTS,  "Phase_2_Raw_MRI_CNN", "figures", "cnn2d_version_comparison.png"),
        os.path.join(PLOTS, "cnn2d_v3", "cnn2d_version_comparison.png"),
    ),
    "fig8_cnn_roc": find_fig(
        os.path.join(OUTS,  "Phase_2_Raw_MRI_CNN", "figures", "cnn2d_v3_roc_curves.png"),
        os.path.join(PLOTS, "cnn2d_v3", "cnn2d_v3_roc_curves.png"),
    ),
    "fig9_gradcam_panel": find_fig(
        os.path.join(OUTS,  "Phase_2_Raw_MRI_CNN", "gradcam", "gradcam_sample_panel.png"),
        os.path.join(PLOTS, "gradcam", "gradcam_sample_panel.png"),
    ),
    "fig10_gradcam_avg": find_fig(
        os.path.join(OUTS,  "Phase_2_Raw_MRI_CNN", "gradcam", "gradcam_average_heatmap.png"),
        os.path.join(PLOTS, "gradcam", "gradcam_average_heatmap.png"),
    ),
    "fig11_fusion_roc": find_fig(
        os.path.join(OUTS,  "Phase_3_Fusion", "figures", "fusion_roc_curves.png"),
        os.path.join(PLOTS, "fusion", "fusion_roc_curves.png"),
    ),
    "fig12_fusion_ablation": find_fig(
        os.path.join(OUTS,  "Phase_3_Fusion", "figures", "fusion_ablation_comparison.png"),
        os.path.join(PLOTS, "fusion", "fusion_ablation_comparison.png"),
    ),
    "fig13_fusion_shap": find_fig(
        os.path.join(OUTS,  "Phase_3_Fusion", "figures", "fusion_shap_analysis.png"),
        os.path.join(PLOTS, "fusion", "fusion_shap_analysis.png"),
    ),
}

print("=" * 60)
print("  Building IEEE Access Paper with Figures")
print("=" * 60)
print("\nFigure availability:")
fig_count = 0
for name, path in FIGURES.items():
    status = f"✓  {path}" if path else "✗  NOT FOUND"
    print(f"  {name:25s}  {status}")
    if path: fig_count += 1
print(f"\n  {fig_count}/{len(FIGURES)} figures found")

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

# ── Helper functions ─────────────────────────────────────────────────────────
fig_counter = [0]

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(11) if level == 1 else Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c) if level == 1 else RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name    = 'Times New Roman'
    run.font.size    = Pt(10)
    run.font.bold    = True
    run.font.italic  = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_para(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    # Handle **bold** inline
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    return p

def add_figure(fig_key, caption_text, width_inches=5.5):
    fig_counter[0] += 1
    n    = fig_counter[0]
    path = FIGURES.get(fig_key)

    if path and os.path.exists(path):
        # Image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        run = p_img.add_run()
        run.add_picture(path, width=Inches(width_inches))

        # Caption
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(f"Fig. {n}. {caption_text}")
        run_cap.font.name   = 'Times New Roman'
        run_cap.font.size   = Pt(9)
        run_cap.font.italic = True
        print(f"  ✓ Embedded Fig. {n}: {os.path.basename(path)}")
    else:
        # Placeholder box
        p_ph = doc.add_paragraph()
        p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ph.paragraph_format.space_before = Pt(8)
        p_ph.paragraph_format.space_after  = Pt(4)
        run_ph = p_ph.add_run(
            f"[ Fig. {n} — {caption_text} — file not found: {fig_key} ]"
        )
        run_ph.font.name   = 'Times New Roman'
        run_ph.font.size   = Pt(9)
        run_ph.font.italic = True
        run_ph.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        print(f"  ⚠ Placeholder Fig. {n}: {fig_key} not found")

def add_table(headers, rows, col_widths_cm, caption_text):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    hdr.height = Cm(0.7)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1a3a5c')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'),   'clear')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name  = 'Times New Roman'
        run.font.size  = Pt(9)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'f0f4ff')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

    # Table caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after  = Pt(12)
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name   = 'Times New Roman'
    run_cap.font.size   = Pt(9)
    run_cap.font.bold   = True
    run_cap.font.italic = True

def add_blank(space_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_pt)

# ═══════════════════════════════════════════════════════════════════════════════
# PAPER CONTENT
# ═══════════════════════════════════════════════════════════════════════════════
print("\nBuilding document...")

# ── TITLE
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(12)
r_title = p_title.add_run(
    "Clinical-First Prediction of MCI-to-Dementia Conversion Using Ensemble Learning, "
    "CNN-Based MRI Analysis, and Multi-Modal Fusion: A Study on ADNI Data"
)
r_title.font.name  = 'Times New Roman'
r_title.font.size  = Pt(16)
r_title.font.bold  = True
r_title.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

for line, sz in [
    ("Gagan [Author Name], Member, IEEE", 12),
    ("Department of Electrical and Electronics Engineering, MIT Manipal,\nManipala Academy of Higher Education, Manipal, India", 10),
    ("Corresponding author: [email address]", 9),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(line)
    r.font.name = 'Times New Roman'; r.font.size = Pt(sz)
    if sz < 11: r.font.italic = True

add_blank(10)

# ── ABSTRACT
p_abs_head = doc.add_paragraph()
r = p_abs_head.add_run("ABSTRACT—")
r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.font.bold = True
r2 = p_abs_head.add_run(
    "Early identification of patients with Mild Cognitive Impairment (MCI) who will progress "
    "to dementia is a critical clinical challenge. This paper presents a systematic multi-phase "
    "pipeline evaluating clinical cognitive assessments, MRI volumetric features, raw MRI deep "
    "learning representations, and their fusion for predicting three-year MCI-to-dementia "
    "conversion. Using 701 subjects from the Alzheimer's Disease Neuroimaging Initiative (ADNI; "
    "273 progressive MCI, 428 stable MCI) and 18 standardised clinical features, an XGBoost "
    "classifier achieved AUC-ROC of 0.821 [0.751, 0.885] with sensitivity of 0.895 at a "
    "screening threshold of 0.15. Adding MRI volumetric features provided no incremental benefit "
    "(AUC = 0.818). A ResNet-18 CNN trained on 14,020 axial brain MRI slices with a "
    "frozen-backbone linear probe achieved AUC = 0.710 [0.606, 0.816] independently. A "
    "comprehensive fusion ablation across 23 strategies failed to statistically improve over "
    "the clinical-only baseline (maximum ΔAUC = +0.004, overlapping 95% CIs). SHAP analysis "
    "identified ADAS-Cog 13, RAVLT Delayed Recall, Digit Span, and CDR Sum of Boxes as "
    "dominant predictors. External clinical validation on OASIS-2 (n = 52) yielded AUC = 0.469, "
    "consistent with cross-cohort feature distribution mismatch. The findings demonstrate that "
    "structured cognitive assessments provide the strongest predictive signal, while MRI features "
    "offer moderate independent but non-additive value."
)
r2.font.name = 'Times New Roman'; r2.font.size = Pt(9)
p_abs_head.paragraph_format.space_after = Pt(4)

p_idx = doc.add_paragraph()
p_idx.paragraph_format.space_after = Pt(16)
ri = p_idx.add_run("INDEX TERMS—")
ri.font.name = 'Times New Roman'; ri.font.size = Pt(9); ri.font.bold = True
ri2 = p_idx.add_run(
    "Mild cognitive impairment, dementia prediction, XGBoost, ResNet-18, convolutional "
    "neural network, multi-modal fusion, SHAP, ADNI, RAVLT, Alzheimer's disease, neuroimaging."
)
ri2.font.name = 'Times New Roman'; ri2.font.size = Pt(9)

# ── I. INTRODUCTION
add_heading("I. Introduction")
add_para(
    "**D**ementia affects over 55 million people worldwide and is projected to reach 139 million "
    "by 2050 [1]. Alzheimer's disease accounts for 60–70% of cases and features a lengthy "
    "preclinical phase during which intervention may be most effective [2]. Mild Cognitive "
    "Impairment (MCI) represents an intermediate state between normal ageing and dementia, "
    "with annual conversion rates of 10–15% [3]. Identifying which MCI patients will convert "
    "within a clinically actionable window is therefore of paramount importance."
)
add_para(
    "Neuropsychological assessments — including the Alzheimer's Disease Assessment Scale "
    "(ADAS-Cog), the Rey Auditory Verbal Learning Test (RAVLT), and the Clinical Dementia "
    "Rating Scale (CDR) — have long served as primary screening tools. Concurrently, advances "
    "in structural neuroimaging and deep learning have motivated efforts to extract predictive "
    "biomarkers directly from brain MRI scans [4]. Whether imaging adds predictive value over "
    "clinical assessments alone remains contested [5], [6]."
)
add_para(
    "This work makes four primary contributions: (i) a systematic phase-by-phase evaluation "
    "of clinical, volumetric MRI, and deep CNN features for pMCI prediction on the full ADNI "
    "cohort; (ii) a comprehensive 23-strategy fusion ablation study; (iii) a SHAP-based "
    "explainability analysis identifying dominant cognitive predictors; and (iv) a clinical "
    "decision-support web application providing accessible risk estimation."
)

# ── II. RELATED WORK
add_heading("II. Related Work")
add_subheading("A. Machine Learning for MCI Conversion Prediction")
add_para(
    "Supervised learning approaches for pMCI prediction have been extensively studied using "
    "ADNI data. SVMs and random forests have achieved AUC values of 0.72–0.82 on clinical "
    "feature sets [7]. XGBoost approaches have demonstrated competitive performance with "
    "improved calibration [8]. The importance of ADAS-Cog, RAVLT, and functional measures "
    "has been consistently reported across studies [9]."
)
add_subheading("B. Deep Learning on MRI for Dementia")
add_para(
    "CNNs applied to structural MRI have shown promise in Alzheimer's disease classification "
    "[10], [11]. MCI conversion prediction presents greater challenges due to subtle structural "
    "changes. ResNet-based architectures on ADNI MRI have reported AUCs of 0.65–0.78 [12]. "
    "Transfer learning from ImageNet has proven effective on small neuroimaging datasets [13]."
)
add_subheading("C. Multi-Modal Fusion")
add_para(
    "Multi-modal fusion combining clinical and imaging features has shown mixed results. "
    "Early concatenation fusion is susceptible to dimensionality imbalance, while late fusion "
    "preserves modality-specific signal [14]. Zhang et al. [15] and Ramírez et al. [16] report "
    "fusion improvements of 2–5% AUC, though these typically involve deep architectural fusion "
    "rather than the feature-level fusion evaluated here."
)

# ── III. METHODS
add_heading("III. Materials and Methods")
add_subheading("A. Dataset")
add_para(
    "Data were obtained from the Alzheimer's Disease Neuroimaging Initiative (ADNI). "
    "The study cohort comprised **701 subjects** with MCI at baseline: **273 progressive MCI "
    "(pMCI)** who converted to dementia within 36 months, and **428 stable MCI (sMCI)** who "
    "did not convert. External validation used 52 subjects from OASIS-2 (14 pMCI-like, "
    "38 sMCI-like)."
)

add_subheading("B. Clinical Feature Set (Phase 1)")
add_para(
    "Eighteen standardised features were extracted from baseline clinical assessments. "
    "Three composite features were derived: MMSE_FAQ_composite = MMSE − FAQ, "
    "ADAS_MMSE_gap = ADAS13 + (30 − MMSE), and RAVLT_forget_rate = "
    "RAVLT_forgetting / (RAVLT_immediate / 5). XGBoost was selected as the primary "
    "classifier with class imbalance handled by scale_pos_weight = 1.57. A 70/15/15 "
    "stratified subject-level split was maintained across all phases."
)

add_table(
    ["Feature", "Description", "Direction"],
    [
        ["ADAS-Cog 13",       "Alzheimer's Disease Assessment Scale (13-item)",   "↑ worse"],
        ["RAVLT Delayed",     "Word list recall after 20–30 min delay",            "↑ better"],
        ["Digit Span",        "Forward + backward digit recall",                   "↑ better"],
        ["CDR Sum of Boxes",  "CDR functional domain sum",                         "↑ worse"],
        ["FAQ",               "Functional Activities Questionnaire",               "↑ worse"],
        ["MMSE",              "Mini-Mental State Examination",                     "↑ better"],
        ["MoCA",              "Montreal Cognitive Assessment",                     "↑ better"],
        ["RAVLT Immediate",   "Cumulative 5-trial word recall",                   "↑ better"],
        ["RAVLT Forgetting",  "Trial 5 minus delayed recall",                     "↓ better"],
        ["Trails B (s)",      "Trail Making Test Part B — time",                  "↓ better"],
        ["ADAS-Cog 11",       "11-item ADAS-Cog",                                 "↑ worse"],
        ["GDS",               "Geriatric Depression Scale",                        "↑ worse"],
        ["CDR Global",        "Global CDR score",                                  "↑ worse"],
        ["Education (yr)",    "Years of formal education",                         "Protective"],
        ["Sex",               "Biological sex (0=M, 1=F)",                        "Covariate"],
        ["MMSE−FAQ",          "Derived: cognitive vs. functional gap",             "Context"],
        ["ADAS+MMSE gap",     "Derived: error-based severity index",              "↑ worse"],
        ["RAVLT forget rate", "Derived: normalised forgetting rate",              "↑ worse"],
    ],
    [4.5, 6.5, 2.0],
    "TABLE I. Clinical Feature Set (18 features used across all phases)"
)

add_subheading("C. MRI Preprocessing (Phase 2)")
add_para(
    "T1-weighted MRI volumes were reoriented to RAS+ canonical orientation, preprocessed "
    "into 2D axial PNG slices (224 × 224 pixels), with 20 slices per subject extracted from "
    "the hippocampal window (axial range: 40–65% of total z-slices). A total of **14,020 slices** "
    "were generated from 701 subjects. Input to the CNN was a three-channel triplet of "
    "consecutive axial slices (N−1, N, N+1) to provide inter-slice context."
)

add_subheading("D. ResNet-18 CNN Architecture")
add_para(
    "Pre-trained ResNet-18 was evaluated in three configurations (Table II). The final "
    "architecture (v3) used a **frozen backbone linear probe**: all convolutional layers frozen, "
    "only the final fully connected layer trained (1,026 parameters). Training used AdamW "
    "(lr = 1×10⁻⁴, weight decay = 1×10⁻⁴) with cosine annealing and class-weighted "
    "cross-entropy. Subject probability was the median over 20 slice-level softmax outputs."
)

add_table(
    ["Version", "Configuration", "Subjects", "Best Epoch", "AUC-ROC", "95% CI"],
    [
        ["v1", "Full fine-tune, 11M trainable params", "442/701", "1 (collapsed)", "0.582", "[0.443, 0.706]"],
        ["v2", "Layer4 + FC, 8.4M trainable", "701/701", "1 (instability)", "0.689", "[0.579, 0.793]"],
        ["v3 ★", "FC-only, 1,026 trainable, triplet input", "701/701", "16 (stable)", "0.710", "[0.606, 0.816]"],
    ],
    [1.0, 4.5, 1.8, 2.2, 1.5, 2.0],
    "TABLE II. CNN Version Comparison — Subject-Level, Median Aggregation"
)

add_subheading("E. Multi-Modal Fusion (Phase 3)")
add_para(
    "CNN embeddings (512-d avgpool features) were compressed to 32 PCA components "
    "(68.0% variance explained, fit on training subjects only). Twenty-three fusion "
    "strategies were evaluated: early feature fusion (XGBoost + LR, PCA ∈ {4,8,16,32}), "
    "late probability weighting (w_clinical ∈ {0.70–0.95}), and uncertain-case correction "
    "(MRI adjusts clinical predictions when p_clinical ∈ [0.35, 0.65])."
)

# ── IV. RESULTS
add_heading("IV. Results")
add_subheading("A. Phase 1: Clinical-Only Performance")
add_para(
    "XGBoost achieved the highest AUC-ROC = **0.821 [0.751, 0.885]** (Table III). "
    "At the screening threshold of 0.15, sensitivity reached **0.895** with specificity "
    "of 0.649. All three models showed substantially higher specificity than sensitivity "
    "at the standard 0.5 threshold, reflecting the class imbalance (pMCI:sMCI = 1:1.57)."
)

add_table(
    ["Model", "AUC-ROC", "95% CI", "F1", "Sens (0.50)", "Spec (0.50)", "Sens (0.15)"],
    [
        ["XGBoost ★",         "0.821", "[0.751, 0.885]", "0.642", "0.614", "0.825", "0.895"],
        ["Random Forest",     "0.798", "[0.722, 0.869]", "0.618", "0.590", "0.811", "—"],
        ["Logistic Regr.",    "0.771", "[0.694, 0.845]", "0.597", "0.571", "0.796", "—"],
    ],
    [3.5, 1.6, 2.2, 1.3, 2.0, 2.0, 2.0],
    "TABLE III. Phase 1 Classifier Comparison (Test Set, n=106). ★ = selected model."
)

add_figure(
    "fig2_shap_beeswarm",
    "SHAP beeswarm plot for XGBoost Phase 1 model. Each dot represents one test subject. "
    "Colour indicates feature value (red=high, blue=low). ADAS-Cog 13 is the dominant predictor.",
    width_inches=5.0
)

add_subheading("B. SHAP Feature Importance")
add_para(
    "Fig. 2 shows the SHAP beeswarm plot. ADAS-Cog 13 contributed the highest mean |SHAP| "
    "value (0.853), followed by RAVLT Delayed Recall (0.625), Digit Span (0.437), and "
    "CDR Sum of Boxes (0.402). These four features collectively explain approximately 67% "
    "of the model's total SHAP signal. Dependence plots (Fig. 3) show a non-linear "
    "relationship for ADAS-Cog 13 with threshold effects near scores of 18–22."
)

add_figure(
    "fig4_shap_dep_adas",
    "SHAP dependence plot for ADAS-Cog 13. Colour indicates RAVLT Delayed Recall interaction. "
    "Risk increases sharply above ADAS-13 ≈ 20, particularly when RAVLT delayed recall is low.",
    width_inches=4.5
)

add_subheading("C. Subgroup Analysis")
add_para(
    "Performance was consistently higher in females (AUC = 0.849 vs. 0.810 in males) "
    "and substantially higher in the high-education subgroup (AUC = 0.871 vs. 0.744). "
    "The 12.7-point education gap is consistent with cognitive reserve theory, wherein "
    "education moderates the relationship between test performance and neuropathology."
)

add_table(
    ["Subgroup", "AUC-ROC"],
    [
        ["Male",              "0.810"],
        ["Female",            "0.849"],
        ["Education ≤ 12 yr", "0.744"],
        ["Education > 12 yr", "0.871"],
        ["Overall",           "0.821"],
    ],
    [6.0, 3.0],
    "TABLE IV. Subgroup AUC-ROC Analysis"
)

add_subheading("D. Phase 1.5: MRI Volumetric Features")
add_para(
    "Adding FreeSurfer MRI volumes did not improve AUC (combined = 0.818; clinical-only = "
    "0.818). SHAP attribution showed clinical features contributed 74.1% and MRI volumes "
    "25.9% of the combined model's explanatory power. MRI volumes alone yielded AUC = 0.638. "
    "Fig. 4 shows the grouped SHAP contribution by modality."
)

add_figure(
    "fig5_phase15_shap",
    "Phase 1.5 SHAP grouped feature importance: clinical features (blue) vs. MRI volumetric "
    "features (orange). Clinical features dominate across all feature importance ranks.",
    width_inches=5.0
)

add_subheading("E. Phase 2: CNN MRI Classification")
add_para(
    "The final ResNet-18 v3 achieved AUC-ROC = **0.710 [0.606, 0.816]** at subject level. "
    "Fig. 5 shows training curves for v3: stable learning over 21 epochs (best epoch = 16) "
    "with train AUC plateauing at ~0.72, confirming no memorisation. Fig. 6 shows the "
    "version comparison across v1, v2, and v3. Fig. 7 presents Grad-CAM activations "
    "showing temporal and parietal focus for pMCI subjects."
)

add_figure(
    "fig6_cnn_curves",
    "CNN v3 training curves. Top: cross-entropy loss. Bottom: AUC-ROC per epoch. "
    "The horizontal dashed line marks the clinical baseline (0.821). Best epoch = 16.",
    width_inches=5.5
)

add_figure(
    "fig7_cnn_compare",
    "CNN version comparison (v1 → v3). AUC improvement from 0.582 to 0.710 achieved "
    "by correcting the subject split and adopting a frozen-backbone linear probe.",
    width_inches=4.0
)

add_figure(
    "fig9_gradcam_panel",
    "Grad-CAM visualisations for representative test subjects. Left: original axial slice. "
    "Right: activation overlay (red/yellow = high attention). pMCI subjects show stronger "
    "activation in temporal-parietal regions consistent with early Alzheimer's atrophy.",
    width_inches=5.5
)

add_figure(
    "fig10_gradcam_avg",
    "Mean Grad-CAM activation maps averaged across all correctly classified test subjects. "
    "Left: pMCI group. Right: sMCI group. Differential activation is concentrated "
    "in hippocampal and entorhinal cortex regions.",
    width_inches=4.5
)

add_subheading("F. Phase 3: Multi-Modal Fusion Ablation")
add_para(
    "Table V presents selected results from the 23-strategy ablation. The clinical-only "
    "baseline (AUC = 0.855 on Phase 3 cohort) was not significantly exceeded by any "
    "strategy. The best result (uncertain-case fusion, ΔAUC = +0.004) had overlapping "
    "95% CIs with the baseline [0.777, 0.920] vs. [0.785, 0.924]. Late fusion strategies "
    "consistently raised sensitivity (+0.025) without meaningful AUC gain, suggesting "
    "MRI provides complementary information specifically for borderline clinical cases. "
    "Figs. 8 and 9 show the ROC comparison and ablation ranking chart."
)

add_table(
    ["Rank", "Strategy", "AUC", "95% CI", "Sens", "Spec"],
    [
        ["1",  "Uncertain-case [0.35,0.65] w_mri=0.3", "0.860", "[0.785, 0.924]", "0.805", "0.738"],
        ["2",  "Early-LR PCA=4",                        "0.859", "[0.782, 0.924]", "0.805", "0.754"],
        ["3",  "Late w=0.75/0.25",                       "0.859", "[0.781, 0.923]", "0.805", "0.754"],
        ["10", "Clinical-only (baseline)",               "0.855", "[0.777, 0.920]", "0.780", "0.815"],
        ["18", "Early-XGB PCA=32",                       "0.822", "[0.742, 0.895]", "0.707", "0.800"],
        ["23", "MRI-only CNN v3",                        "0.710", "[0.604, 0.817]", "0.854", "0.338"],
    ],
    [0.7, 4.8, 1.2, 2.2, 1.2, 1.2],
    "TABLE V. Fusion Ablation — Selected Strategies (Test Set, n=106; Phase 3 cohort)"
)

add_figure(
    "fig11_fusion_roc",
    "ROC curves for Phase 3 fusion strategies. Clinical-only (blue solid), "
    "late fusion w=0.75/0.25 (green dashed), MRI-only CNN (orange). "
    "Confidence bands (shaded) overlap for all strategies vs. clinical baseline.",
    width_inches=4.5
)

add_figure(
    "fig12_fusion_ablation",
    "Ranked AUC of all 23 fusion strategies with 95% bootstrap confidence intervals. "
    "The clinical-only baseline (rank 10, blue) is not significantly outperformed. "
    "Late fusion and uncertain-case strategies (green) cluster near the baseline.",
    width_inches=5.5
)

add_figure(
    "fig13_fusion_shap",
    "SHAP feature group contribution in the fusion model. Left: % of total |SHAP| "
    "by modality (clinical 50.9%, MRI-PCA 49.1%). Right: top 15 features coloured "
    "by modality — ADAS-Cog 13 remains the top feature, followed by MRI_PC1 and MRI_PC4.",
    width_inches=5.5
)

add_subheading("G. External Validation (OASIS-2)")
add_para(
    "External clinical validation on OASIS-2 (n = 52) yielded AUC = 0.469 [0.312, 0.619] — "
    "effectively at chance. This is attributed to systematic feature distribution mismatch "
    "across cohorts (different RAVLT and ADAS protocols, age distributions, label definitions) "
    "rather than model failure. Only 3 of 18 ADNI features were directly comparable in OASIS-2."
)

# ── V. DISCUSSION
add_heading("V. Discussion")
add_para(
    "The principal finding of this study is that structured cognitive and functional assessments "
    "provide robust, reproducible, and sufficient signal for MCI-to-dementia conversion "
    "prediction, with no fusion strategy offering statistically significant improvement. "
    "This is consistent with earlier work by Moradi et al. [18] and Zhang et al. [19], who "
    "reported that clinical features dominate imaging-derived features in cross-validated pMCI "
    "prediction."
)
add_para(
    "The failure of CNN embedding fusion is informative. Despite training on a meaningful "
    "cohort (14,020 slices, 701 subjects), the CNN achieved only moderate subject-level "
    "AUC (0.710), reflecting the inherent difficulty of detecting MCI-stage structural "
    "changes in 2D axial slices. PCA-compressed embeddings carry sufficient noise to "
    "dilute clinical signal in early concatenation fusion. The modest late-fusion "
    "sensitivity gain (+0.025) is clinically relevant in screening contexts but insufficient "
    "to justify added MRI processing complexity in routine practice."
)
add_para(
    "The 12.7-percentage-point education gap in subgroup AUC (0.871 high vs. 0.744 low "
    "education) reflects the documented education bias in cognitive testing. Future work "
    "should evaluate education-adjusted normative scoring or stratified models."
)
add_para(
    "**Limitations.** First, all models were trained on a single ADNI cohort; cross-cohort "
    "generalisation is not yet established. Second, 2D CNNs do not capture full 3D volumetric "
    "information and may underperform against 3D architectures on larger datasets. Third, "
    "temporal dynamics (rate of cognitive decline) were not modelled. Fourth, the frozen "
    "backbone may limit MRI-specific feature adaptation. Future work should explore "
    "longitudinal modelling, attention-based architectures, and multi-site training."
)

# ── VI. CONCLUSION
add_heading("VI. Conclusion")
add_para(
    "This paper presented a systematic multi-phase evaluation of clinical and imaging predictors "
    "for MCI-to-dementia conversion using 701 ADNI subjects. The clinical-only XGBoost model "
    "achieved AUC = 0.821 with sensitivity of 0.895 in screening mode — sufficient for "
    "practical clinical screening. MRI volumetric features and CNN embeddings both provided "
    "independent signal (AUC 0.638 and 0.710 respectively) but offered no statistically "
    "significant improvement across 23 tested fusion strategies. SHAP analysis confirmed "
    "ADAS-Cog 13, RAVLT Delayed Recall, Digit Span, and CDR Sum of Boxes as the dominant "
    "predictive features. These results suggest that investment in standardised, high-quality "
    "cognitive assessment — rather than additional imaging — provides the most cost-effective "
    "path to early dementia risk stratification in MCI populations."
)

# ── ACKNOWLEDGMENT
add_heading("Acknowledgment")
add_para(
    "Data collection and sharing for this project was funded by the Alzheimer's Disease "
    "Neuroimaging Initiative (ADNI) (National Institutes of Health Grant U01 AG024904). "
    "ADNI is funded by the National Institute on Aging, the National Institute of Biomedical "
    "Imaging and Bioengineering, and through contributions from numerous pharmaceutical "
    "companies and the non-profit organisation Foundation for the National Institutes of "
    "Health (www.fnih.org). The authors declare no conflicts of interest."
)

# ── REFERENCES
add_heading("References")
refs = [
    "World Health Organization, \"Dementia,\" WHO Fact Sheet, 2023.",
    "C. R. Jack et al., \"NIA-AA Research Framework: Toward a biological definition of Alzheimer's disease,\" Alzheimer's & Dementia, vol. 14, no. 4, pp. 535–562, 2018.",
    "R. C. Petersen et al., \"Mild cognitive impairment: Ten years later,\" Archives of Neurology, vol. 66, no. 12, pp. 1447–1455, 2009.",
    "S. Rathore et al., \"A review on neuroimaging-based classification studies for Alzheimer's disease and its prodromal stages,\" NeuroImage, vol. 155, pp. 530–548, 2017.",
    "B. Fischl, \"FreeSurfer,\" NeuroImage, vol. 62, no. 2, pp. 774–781, 2012.",
    "E. Westman et al., \"Combining MRI and CSF measures for classification of Alzheimer's disease,\" NeuroImage, vol. 62, no. 1, pp. 229–238, 2012.",
    "S. Klöppel et al., \"Automatic classification of MR scans in Alzheimer's disease,\" Brain, vol. 131, no. 3, pp. 681–689, 2008.",
    "T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" in Proc. ACM SIGKDD, 2016, pp. 785–794.",
    "A. Eshaghi et al., \"Progression of regional grey matter atrophy,\" Brain, vol. 141, no. 6, pp. 1665–1677, 2018.",
    "J. Islam and Y. Zhang, \"Brain MRI analysis for Alzheimer's disease diagnosis using deep CNNs,\" Brain Informatics, vol. 5, no. 2, p. 2, 2018.",
    "P. Bhagwat et al., \"Modeling clinical symptom trajectories in Alzheimer's disease,\" PLOS Computational Biology, vol. 14, no. 9, 2018.",
    "K. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning for image recognition,\" in Proc. IEEE CVPR, 2016, pp. 770–778.",
    "A. Ravi et al., \"Effective use of pretrained deep learning models in clinical image analysis,\" in Proc. MICCAI, 2019.",
    "D. Ramírez et al., \"Ensemble of random forests classifiers for MCI and AD prediction,\" J. Neuroscience Methods, vol. 302, pp. 47–57, 2018.",
    "D. Zhang et al., \"Multimodal classification of Alzheimer's disease and MCI,\" NeuroImage, vol. 55, no. 3, pp. 856–867, 2011.",
    "J. Ramírez et al., \"Multimodal classification of Alzheimer's disease,\" in Proc. IEEE EMBC, 2016.",
    "S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" in Advances in NeurIPS, 2017.",
    "E. Moradi et al., \"Machine learning framework for MRI-based Alzheimer's conversion prediction,\" NeuroImage, vol. 104, pp. 398–412, 2015.",
    "D. Zhang et al., \"Predicting future clinical changes of MCI patients,\" PLOS ONE, vol. 6, no. 3, 2011.",
]
for i, ref in enumerate(refs, 1):
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_after  = Pt(2)
    p_ref.paragraph_format.first_line_indent = Cm(-0.5)
    p_ref.paragraph_format.left_indent  = Cm(0.5)
    r_num = p_ref.add_run(f"[{i}] ")
    r_num.font.name = 'Times New Roman'; r_num.font.size = Pt(9); r_num.font.bold = True
    r_txt = p_ref.add_run(ref)
    r_txt.font.name = 'Times New Roman'; r_txt.font.size = Pt(9)

# ── SAVE
doc.save(OUT_FILE)
print(f"\n{'='*60}")
print(f"  COMPLETE")
print(f"{'='*60}")
print(f"  Saved: {OUT_FILE}")
print(f"  Figures embedded: {fig_counter[0]}")
print(f"\n  Next steps:")
print(f"  1. Add your full name and email in the author block")
print(f"  2. Check IEEE Access formatting guide at ieee.org/access")
print(f"  3. Add missing figures if any showed ⚠ above")